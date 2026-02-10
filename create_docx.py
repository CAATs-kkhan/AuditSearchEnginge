"""Create Word document for Audit Document Search Engine"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Audit Document Search Engine', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('A powerful search tool for auditors with Named Entity Recognition (NER), Topic Modelling, Sentiment Analysis, and Advanced AI Features.')

# Features Section
doc.add_heading('Features', level=1)

doc.add_heading('Core Search Capabilities', level=2)
doc.add_paragraph('Simple Search - Basic keyword search across all documents', style='List Bullet')
doc.add_paragraph('Boolean Search - Support for AND, OR, NOT operators', style='List Bullet')
doc.add_paragraph('Wildcard Search - Use * for pattern matching (e.g., procur*)', style='List Bullet')
doc.add_paragraph('Fuzzy Search - Find partial matches and similar terms', style='List Bullet')

doc.add_heading('Document Support', level=2)
doc.add_paragraph('PDF files (with OCR for scanned documents)', style='List Bullet')
doc.add_paragraph('Word documents (.docx, .doc)', style='List Bullet')
doc.add_paragraph('Excel spreadsheets (.xlsx, .xls)', style='List Bullet')
doc.add_paragraph('Text files (.txt, .md)', style='List Bullet')

doc.add_heading('Named Entity Recognition (NER)', level=2)
doc.add_paragraph('Persons - Names and titles', style='List Bullet')
doc.add_paragraph('Organizations - Companies, government bodies, institutions', style='List Bullet')
doc.add_paragraph('Locations - Cities, countries (Pakistani and international)', style='List Bullet')
doc.add_paragraph('Dates - Various date formats', style='List Bullet')
doc.add_paragraph('Money - Currency amounts (USD, PKR, Rs., etc.)', style='List Bullet')
doc.add_paragraph('Emails and Phone Numbers', style='List Bullet')
doc.add_paragraph('Percentages', style='List Bullet')

doc.add_heading('PDF Table Extraction', level=2)
doc.add_paragraph('Extract tables from multiple PDFs', style='List Bullet')
doc.add_paragraph('Export to consolidated Excel file', style='List Bullet')
doc.add_paragraph('Preview extracted tables in the app', style='List Bullet')

doc.add_heading('Topic Modelling (LDA)', level=2)
doc.add_paragraph('Automatically discover themes in documents', style='List Bullet')
doc.add_paragraph('Assign documents to topics', style='List Bullet')
doc.add_paragraph('Filter documents by topic', style='List Bullet')
doc.add_paragraph('Configurable number of topics and words', style='List Bullet')

doc.add_heading('Sentiment and Risk Analysis', level=2)
doc.add_paragraph('Positive/Neutral/Negative sentiment detection', style='List Bullet')
doc.add_paragraph('Risk level assessment (High/Medium/Low)', style='List Bullet')
doc.add_paragraph('Risk keyword detection:', style='List Bullet')
doc.add_paragraph('    - High Risk: fraud, embezzlement, violation, corruption, theft')
doc.add_paragraph('    - Medium Risk: delay, overdue, pending, error, deviation')
doc.add_paragraph('    - Attention: urgent, critical, priority, investigate')

doc.add_heading('AI Assistant (Phase 9)', level=2)

doc.add_heading('Document Summarization', level=3)
doc.add_paragraph('Auto-generate summaries using T5 model', style='List Bullet')
doc.add_paragraph('Summaries saved for quick access', style='List Bullet')
doc.add_paragraph('Works with any indexed document', style='List Bullet')

doc.add_heading('Q&A Chat (RAG)', level=3)
doc.add_paragraph('Ask questions in natural language', style='List Bullet')
doc.add_paragraph('Retrieves relevant document chunks', style='List Bullet')
doc.add_paragraph('Uses FAISS for fast vector search', style='List Bullet')
doc.add_paragraph('Sentence-transformers for embeddings', style='List Bullet')

doc.add_heading('Anomaly Detection', level=3)
doc.add_paragraph('Extracts monetary amounts from documents', style='List Bullet')
doc.add_paragraph('Uses Isolation Forest algorithm', style='List Bullet')
doc.add_paragraph('Flags statistical outliers', style='List Bullet')
doc.add_paragraph('Supports multiple currency formats', style='List Bullet')

# Installation Section
doc.add_page_break()
doc.add_heading('Installation', level=1)

doc.add_heading('Prerequisites', level=2)
doc.add_paragraph('Python 3.10+', style='List Bullet')
doc.add_paragraph('pip package manager', style='List Bullet')

doc.add_heading('Install Dependencies', level=2)

doc.add_paragraph('Core dependencies:')
p = doc.add_paragraph('pip install streamlit pypdf python-docx openpyxl')
p.runs[0].font.name = 'Courier New'

doc.add_paragraph('NLP dependencies:')
p = doc.add_paragraph('pip install spacy textblob scikit-learn')
p.runs[0].font.name = 'Courier New'

doc.add_paragraph('OCR dependencies (optional):')
p = doc.add_paragraph('pip install pytesseract pdf2image pillow')
p.runs[0].font.name = 'Courier New'

doc.add_paragraph('Table extraction:')
p = doc.add_paragraph('pip install pdfplumber')
p.runs[0].font.name = 'Courier New'

doc.add_paragraph('AI Features (Phase 9):')
p = doc.add_paragraph('pip install transformers torch sentence-transformers faiss-cpu')
p.runs[0].font.name = 'Courier New'

# Usage Section
doc.add_heading('Usage', level=1)

doc.add_heading('Running the Application', level=2)
p = doc.add_paragraph('cd AuditSearchEngine')
p.runs[0].font.name = 'Courier New'
p = doc.add_paragraph('streamlit run app/search_app.py')
p.runs[0].font.name = 'Courier New'
doc.add_paragraph('The app will open at http://localhost:8501')

doc.add_heading('Indexing Documents', level=2)
doc.add_paragraph('1. Place documents in the documents/ folder')
doc.add_paragraph('2. Click "Re-index with NER" in the sidebar')
doc.add_paragraph('3. Wait for indexing to complete')

doc.add_heading('Using Search', level=2)
doc.add_paragraph('1. Go to the Search tab')
doc.add_paragraph('2. Enter your query')
doc.add_paragraph('3. Select search type (Simple, Boolean, Wildcard, or Fuzzy)')
doc.add_paragraph('4. View results with highlighted matches')
doc.add_paragraph('5. Export to Excel if needed')

doc.add_heading('Using AI Features', level=2)

doc.add_heading('Summarization', level=3)
doc.add_paragraph('1. Go to AI Assistant > Summarization')
doc.add_paragraph('2. Select a document')
doc.add_paragraph('3. Click "Generate Summary"')
doc.add_paragraph('4. View and save the summary')

doc.add_heading('Q&A Chat', level=3)
doc.add_paragraph('1. Go to AI Assistant > Q&A Chat')
doc.add_paragraph('2. Click "Index Embeddings" (first time only)')
doc.add_paragraph('3. Type your question')
doc.add_paragraph('4. View relevant document excerpts')

doc.add_heading('Anomaly Detection', level=3)
doc.add_paragraph('1. Go to AI Assistant > Anomaly Detection')
doc.add_paragraph('2. Click "Run Anomaly Detection"')
doc.add_paragraph('3. Review flagged amounts')

# Project Structure
doc.add_page_break()
doc.add_heading('Project Structure', level=1)
structure = """AuditSearchEngine/
    app/
        search_app.py          # Main application
    documents/                 # Place documents here
    search_index.json          # Document index
    entities_index.json        # Extracted entities
    topics_index.json          # Topic model results
    sentiment_index.json       # Sentiment analysis results
    summaries_index.json       # Generated summaries
    faiss_index.bin            # Vector embeddings index
    faiss_data.json            # Chunk data for FAISS"""
p = doc.add_paragraph(structure)
p.runs[0].font.name = 'Courier New'
p.runs[0].font.size = Pt(9)

# Tabs Overview
doc.add_heading('Tabs Overview', level=1)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Tab'
hdr_cells[1].text = 'Description'
# Make header bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data = [
    ('Search', 'Main search interface with multiple search modes'),
    ('Entities (NER)', 'View extracted persons, organizations, locations, etc.'),
    ('Statistics', 'Document statistics and word counts'),
    ('Table Extraction', 'Extract and export tables from PDFs'),
    ('Topics', 'LDA topic modelling results'),
    ('Sentiment', 'Risk and sentiment analysis'),
    ('AI Assistant', 'Summarization, Q&A, and Anomaly Detection'),
]
for i, (tab, desc) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = tab
    row[1].text = desc

# Troubleshooting
doc.add_heading('Troubleshooting', level=1)

doc.add_heading('"Summarization not available"', level=2)
doc.add_paragraph('Install transformers and torch:')
p = doc.add_paragraph('pip install transformers torch')
p.runs[0].font.name = 'Courier New'

doc.add_heading('"Q&A Chat not available"', level=2)
doc.add_paragraph('Install sentence-transformers and faiss:')
p = doc.add_paragraph('pip install sentence-transformers faiss-cpu')
p.runs[0].font.name = 'Courier New'

doc.add_heading('OCR not working', level=2)
doc.add_paragraph('1. Verify Tesseract is installed')
doc.add_paragraph('2. Check Poppler path in C:\\poppler\\')
doc.add_paragraph('3. Ensure pdf2image is installed')

# Version History
doc.add_heading('Version History', level=1)
doc.add_paragraph('Phase 1-3: Core search, NER, document support', style='List Bullet')
doc.add_paragraph('Phase 4: PDF table extraction', style='List Bullet')
doc.add_paragraph('Phase 5: Topic modelling (LDA)', style='List Bullet')
doc.add_paragraph('Phase 6: Sentiment and risk analysis', style='List Bullet')
doc.add_paragraph('Phase 7: OCR support for scanned PDFs', style='List Bullet')
doc.add_paragraph('Phase 9: AI features (Summarization, Q&A, Anomaly Detection)', style='List Bullet')

# Footer
doc.add_page_break()
doc.add_heading('License', level=1)
doc.add_paragraph('Internal audit tool - For authorized use only.')

doc.add_heading('Support', level=1)
doc.add_paragraph('For issues or feature requests, contact the development team.')

# Save
doc.save('C:/Users/user/Documents/code/AI_P900/Software_Development/AuditSearchEngine/Audit_Document_Search.docx')
print('Document created successfully: Audit_Document_Search.docx')
